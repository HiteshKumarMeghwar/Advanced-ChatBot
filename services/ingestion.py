from db.models import Document, DocumentChunk
from sqlalchemy.ext.asyncio import AsyncSession
from services.vector_db_faiss import FAISSVectorDB
from core.config import MAX_CHUNK_CHARS, MIME_MAP
from langchain_core.documents import Document

import fitz                                 # PyMuPDF
import docx
from pathlib import Path
from typing import List
from unstructured.partition.auto import partition
from unstructured.chunking.title import chunk_by_title
import aiofiles
import asyncio
import nltk
try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    nltk.download('punkt_tab', quiet=True)



class DocumentIngestor:
    def __init__(self, db: AsyncSession):
        self.db = db

    async def ingest_document(
        self,
        document_id: int,
        document_path: str,
        document_name: str,
        thread_id: str,
        vector_db: FAISSVectorDB,
    ):
        # Extract chunks asynchronously
        text_chunks = await self._extract_chunks(document_path)

        # persist chunks in MySQL
        db_chunks = [
            DocumentChunk(
                document_id=document_id,
                chunk_index=idx,
                text=chunk
            )
            for idx, chunk in enumerate(text_chunks)
        ]
        self.db.add_all(db_chunks)
        await self.db.flush() 

        # LangChain Document objects
        lc_docs = [
            Document(
                page_content=chunk,
                metadata={
                    "document_id": str(document_id),
                    "chunk_index": str(idx),
                    "chunk_id": str(db_chunks[idx].id),
                    "source": document_name
                }
            )
            for idx, chunk in enumerate(text_chunks)
        ]

        # store in FAISS (thread isolated)
        await vector_db.add_documents(
            thread_id=str(thread_id),
            documents=lc_docs,
            db=self.db
        )
        
        await self.db.commit()
        # return [c.text for c in db_chunks]        # DTO friendly
        return len(db_chunks)
    

    def _split_text(self, text: str) -> List[str]:
        """
        Split long text into <= MAX_CHUNK_CHARS chunks, preferably at sentence
        boundary.  Uses nltk.sent_tokenize for robustness.
        """
        import nltk
        try:
            sentences = nltk.sent_tokenize(text)
        except LookupError:                      # first run only
            nltk.download('punkt')
            sentences = nltk.sent_tokenize(text)

        chunks, current = [], ""
        for sent in sentences:
            if len(current) + len(sent) > MAX_CHUNK_CHARS:
                if current:
                    chunks.append(current.strip())
                    current = sent
                else:                            # sentence itself too long
                    chunks.append(sent[:MAX_CHUNK_CHARS])
            else:
                current += " " + sent
        if current:
            chunks.append(current.strip())
        return chunks


    # ------------------------------------------------------------------
    # production extractor
    # ------------------------------------------------------------------
    async def _extract_chunks(self, file_path: str):
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._extract_chunks_sync, file_path)
    
    def _extract_chunks_sync(self, file_path: str) -> List[str]:
        """
        Production-grade text extractor.
        Returns clean, deduplicated, length-aware chunks.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(path)

        ext = path.suffix.lower()
        doc_type = MIME_MAP.get(ext, "unstructured")

        all_text: List[str] = []

        # ---------------- PDF ----------------
        if doc_type == "pdf":
            doc = fitz.open(path)
            for page in doc:
                text = page.get_text("text")
                if text.strip():
                    all_text.extend(self._split_text(text))
            doc.close()

        # ---------------- DOCX ----------------
        elif doc_type == "docx":
            doc = docx.Document(path)
            full_text = []
            for para in doc.paragraphs:
                txt = para.text.strip()
                if txt:
                    full_text.append(txt)
            text = "\n".join(full_text)
            all_text.extend(self._split_text(text))

        # ---------------- everything else (unstructured) ----------------
        else:
            elements = partition(file_path, strategy="fast")
            chunks = chunk_by_title(elements, max_characters=MAX_CHUNK_CHARS)
            all_text = [c.text.strip() for c in chunks if c.text.strip()]

        # ---- deduplicate & final length check ----
        seen = set()
        final = []
        for chunk in all_text:
            if chunk not in seen and chunk:
                seen.add(chunk)
                final.append(chunk)
        return final
